from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO

document = Document()

document.add_heading("簡単なWordドキュメントのタイトル", 0)
document.add_paragraph("簡単なWordドキュメントのテキスト")

document.save("sample.docx")


# 1. 「sample.docx」を読み込んで画像を貼り付ける
def insert_image_from_url(docx_file, image_url):
    doc = Document(docx_file)
    response = requests.get(image_url)
    image_stream = BytesIO(response.content)
    doc.add_picture(
        image_stream, width=Inches(3)
    )  # 画像を挿入する。適切なサイズを調整してください。
    doc.save("sample_answer.docx")  # 変更を保存する


# 2. 「sample.docx」内の文字数をカウントして出力する
def count_characters(docx_file):
    doc = Document(docx_file)
    total_characters = 0
    for paragraph in doc.paragraphs:
        total_characters += len(paragraph.text)
    print("Total characters in the document:", total_characters)


# 3. 「sample_answer.docx」としてドキュメントを保存する
def save_document_as_new(docx_file, new_file_name):
    doc = Document(docx_file)
    doc.save(new_file_name)


if __name__ == "__main__":
    # 1. 画像をURLから取得して挿入する
    image_url = "https://i.pinimg.com/originals/fd/50/47/fd5047c145bdcc643b4b84c67d219ddf.jpg"  # 画像のURLを指定する
    insert_image_from_url("sample.docx", image_url)

    # 2. 文字数をカウントして出力する
    count_characters("sample.docx")
